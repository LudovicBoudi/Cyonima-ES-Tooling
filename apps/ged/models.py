from django.db import models, transaction
from django.contrib.auth.models import User
from django.utils import timezone
import os
import io
import re
import uuid
import logging

logger = logging.getLogger(__name__)


def extract_text(file_field, file_type):
    content = file_field.read()
    file_field.seek(0)
    text = ''
    if file_type == 'txt':
        text = content.decode('utf-8', errors='replace')
    elif file_type == 'pdf':
        from pdfminer.high_level import extract_text as pdf_extract
        text = pdf_extract(io.BytesIO(content))
    elif file_type == 'docx':
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        text = '\n'.join(p.text for p in doc.paragraphs)
    text = re.sub(r'\s+', ' ', text).strip()
    return text[:50000]  # limit to 50k chars


def next_registration_number():
    year = timezone.now().year
    prefix = f'DOC-{year}-'
    with transaction.atomic():
        last = Document.objects.select_for_update().filter(
            registration_number__startswith=prefix, deleted_at__isnull=True
        ).order_by('registration_number').last()
        if last:
            seq = int(last.registration_number.split('-')[-1]) + 1
        else:
            seq = 1
        return f'{prefix}{seq:05d}'


class DocumentCategory(models.Model):
    name = models.CharField('Nom', max_length=100)
    color = models.CharField('Couleur', max_length=7, default='#1a3a6b', help_text='Hex (#1a3a6b)')
    created_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, editable=False)

    class Meta:
        verbose_name = 'Catégorie'
        verbose_name_plural = 'Catégories'
        ordering = ['name']

    def __str__(self):
        return self.name


def document_upload_path(instance, filename):
    return f'ged/{instance.category.name if instance.category else "divers"}/{filename}'


class Document(models.Model):
    STATUS_CHOICES = [
        ('brouillon', 'Brouillon'),
        ('en_relecture', 'En relecture'),
        ('publie', 'Publié'),
        ('archive', 'Archivé'),
    ]
    registration_number = models.CharField('N° enregistrement', max_length=20, unique=True, editable=False, null=True)
    title = models.CharField('Titre', max_length=255)
    file = models.FileField('Fichier', upload_to=document_upload_path)
    category = models.ForeignKey(DocumentCategory, on_delete=models.SET_NULL, null=True, blank=True, related_name='documents', verbose_name='Catégorie')
    status = models.CharField('Statut', max_length=20, choices=STATUS_CHOICES, default='brouillon')
    description = models.TextField('Description', blank=True)
    version = models.CharField('Version', max_length=20, default='1.0')
    tags = models.CharField('Tags', max_length=500, blank=True, help_text='Séparés par des virgules')
    file_size = models.BigIntegerField('Taille (octets)', editable=False, default=0)
    file_type = models.CharField('Type', max_length=100, editable=False, blank=True)
    download_count = models.PositiveIntegerField('Téléchargements', default=0, editable=False)
    content_text = models.TextField('Texte extrait', blank=True, editable=False)
    reviewed_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, blank=True, related_name='ged_reviewed')
    reviewed_at = models.DateTimeField(null=True, blank=True)
    created_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, editable=False)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    deleted_at = models.DateTimeField(null=True, blank=True, verbose_name='Supprimé le')

    class Meta:
        verbose_name = 'Document'
        verbose_name_plural = 'Documents'
        ordering = ['-updated_at']

    def save(self, *args, **kwargs):
        if not self.registration_number:
            self.registration_number = next_registration_number()
        if self.file:
            self.file_size = self.file.size
            _, ext = os.path.splitext(self.file.name)
            self.file_type = ext.lower().lstrip('.') if ext else 'inconnu'
        if self.file and self.file_type in ('txt', 'pdf', 'docx'):
            try:
                self.content_text = extract_text(self.file, self.file_type)
            except Exception as e:
                logger.warning(f"Échec extraction texte pour {self.title}: {e}")
        super().save(*args, **kwargs)

    def filename(self):
        return os.path.basename(self.file.name) if self.file else ''

    def size_display(self):
        s = self.file_size
        if s >= 1024 * 1024:
            return f'{s / (1024 * 1024):.1f} Mo'
        elif s >= 1024:
            return f'{s / 1024:.1f} Ko'
        return f'{s} o'

    def __str__(self):
        return f'{self.title} (v{self.version})'


class DocumentVersion(models.Model):
    document = models.ForeignKey(Document, on_delete=models.CASCADE, related_name='versions')
    file = models.FileField('Fichier', upload_to='ged/versions/%Y/%m/')
    version_number = models.CharField('Version', max_length=20)
    file_size = models.BigIntegerField('Taille (octets)', editable=False, default=0)
    file_type = models.CharField('Type', max_length=100, editable=False, blank=True)
    notes = models.CharField('Notes', max_length=500, blank=True)
    uploaded_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True)
    uploaded_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = 'Version'
        verbose_name_plural = 'Versions'
        ordering = ['-uploaded_at']

    def save(self, *args, **kwargs):
        if self.file:
            self.file_size = self.file.size
            _, ext = os.path.splitext(self.file.name)
            self.file_type = ext.lower().lstrip('.') if ext else 'inconnu'
        super().save(*args, **kwargs)

    def filename(self):
        return os.path.basename(self.file.name) if self.file else ''

    def size_display(self):
        s = self.file_size
        if s >= 1024 * 1024:
            return f'{s / (1024 * 1024):.1f} Mo'
        elif s >= 1024:
            return f'{s / 1024:.1f} Ko'
        return f'{s} o'

    def __str__(self):
        return f'{self.document.title} v{self.version_number}'


class SharedLink(models.Model):
    document = models.ForeignKey(Document, on_delete=models.CASCADE, related_name='shared_links')
    token = models.UUIDField('Token', unique=True, default=uuid.uuid4, editable=False)
    expires_at = models.DateTimeField('Expire le')
    is_active = models.BooleanField('Actif', default=True)
    created_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True)
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = 'Lien de partage'
        verbose_name_plural = 'Liens de partage'
        ordering = ['-created_at']

    def is_expired(self):
        return self.expires_at < timezone.now()

    def __str__(self):
        return f'Partage {self.document.title} — {"actif" if not self.is_expired() and self.is_active else "expiré/inactif"}'


class UserFavorite(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='ged_favorites')
    document = models.ForeignKey(Document, on_delete=models.CASCADE, related_name='favorited_by')
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = 'Favori'
        verbose_name_plural = 'Favoris'
        unique_together = ('user', 'document')
        ordering = ['-created_at']

    def __str__(self):
        return f'{self.user.username} → {self.document.title}'


class CategorySubscription(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='ged_subscriptions')
    category = models.ForeignKey(DocumentCategory, on_delete=models.CASCADE, related_name='subscribers')
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = 'Abonnement'
        verbose_name_plural = 'Abonnements'
        unique_together = ('user', 'category')
        ordering = ['-created_at']

    def __str__(self):
        return f'{self.user.username} → {self.category.name}'


class AuditLog(models.Model):
    ACTION_CHOICES = [
        ('create', 'Création'),
        ('edit', 'Modification'),
        ('delete', 'Suppression (corbeille)'),
        ('restore', 'Restauration'),
        ('permanent_delete', 'Suppression définitive'),
        ('submit', 'Soumission en relecture'),
        ('approve', 'Approbation / Publication'),
        ('archive', 'Archivage'),
        ('unarchive', 'Désarchivage'),
        ('version_restore', 'Restauration de version'),
        ('share', 'Partage par lien'),
        ('download', 'Téléchargement'),
    ]
    document = models.ForeignKey(Document, on_delete=models.CASCADE, related_name='audit_logs')
    user = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, blank=True)
    action = models.CharField('Action', max_length=30, choices=ACTION_CHOICES)
    details = models.TextField('Détails', blank=True)
    ip_address = models.GenericIPAddressField('IP', blank=True, null=True)
    created_at = models.DateTimeField('Date', auto_now_add=True)

    class Meta:
        verbose_name = "Journal d'audit"
        verbose_name_plural = "Journaux d'audit"
        ordering = ['-created_at']

    def __str__(self):
        return f'{self.get_action_display()} — {self.document.title} ({self.created_at.strftime("%d/%m/%Y %H:%M")})'
